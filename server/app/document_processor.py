import os
import re
from collections import Counter
from typing import Dict, List, Optional

from docx import Document as DocxDocument
from pypdf import PdfReader


STOPWORDS = {
    "the", "and", "for", "with", "from", "this", "that", "shall", "will",
    "are", "was", "were", "you", "your", "can", "not", "all", "any",
    "into", "then", "than", "have", "has", "had", "but", "or", "if",
    "in", "on", "of", "to", "a", "an", "is", "be", "by", "as", "at",
    "it", "its", "may", "must", "should"
}


def clean_text(text: str) -> str:
    """
    Normalize whitespace and remove excessive line breaks.
    """
    if not text:
        return ""

    text = text.replace("\x00", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = text.strip()

    return text


def extract_text_from_txt(file_path: str) -> List[Dict]:
    """
    Extract text from TXT or MD files.
    Returns a list of page-like records for consistent handling.
    """
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        text = f.read()

    return [
        {
            "page_number": None,
            "text": clean_text(text)
        }
    ]


def extract_text_from_pdf(file_path: str) -> List[Dict]:
    """
    Extract text from each page of a PDF.
    """
    reader = PdfReader(file_path)
    pages = []

    for idx, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        text = clean_text(text)

        if text:
            pages.append(
                {
                    "page_number": idx + 1,
                    "text": text
                }
            )

    return pages


def extract_text_from_docx(file_path: str) -> List[Dict]:
    """
    Extract text from DOCX paragraphs and tables.
    """
    doc = DocxDocument(file_path)
    parts = []

    for para in doc.paragraphs:
        if para.text and para.text.strip():
            parts.append(para.text.strip())

    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                cell_text = clean_text(cell.text)
                if cell_text:
                    row_text.append(cell_text)
            if row_text:
                parts.append(" | ".join(row_text))

    text = clean_text("\n".join(parts))

    return [
        {
            "page_number": None,
            "text": text
        }
    ]


def extract_document_text(file_path: str) -> List[Dict]:
    """
    Detect file type and extract text.
    """
    ext = os.path.splitext(file_path)[1].lower()

    if ext in [".txt", ".md"]:
        return extract_text_from_txt(file_path)

    if ext == ".pdf":
        return extract_text_from_pdf(file_path)

    if ext == ".docx":
        return extract_text_from_docx(file_path)

    raise ValueError(f"Unsupported file type for text extraction: {ext}")


def chunk_text(text: str, chunk_size_words: int = 140, overlap_words: int = 30) -> List[str]:
    """
    Split text into overlapping word chunks.

    Example:
    chunk_size_words = 140
    overlap_words = 30

    This means each chunk contains around 140 words and the next chunk
    repeats the last 30 words from the previous chunk.
    """
    text = clean_text(text)

    if not text:
        return []

    words = text.split()

    if len(words) <= chunk_size_words:
        return [" ".join(words)]

    chunks = []
    start = 0

    while start < len(words):
        end = start + chunk_size_words
        chunk = words[start:end]

        if chunk:
            chunks.append(" ".join(chunk))

        if end >= len(words):
            break

        start = end - overlap_words

        if start < 0:
            start = 0

    return chunks


def tokenize(text: str) -> List[str]:
    """
    Convert text into lowercase searchable tokens.
    """
    text = text.lower()
    tokens = re.findall(r"[a-zA-Z0-9_\-/\.]+", text)

    return [
        token for token in tokens
        if len(token) >= 2 and token not in STOPWORDS
    ]


def generate_keywords(text: str, max_keywords: int = 25) -> str:
    """
    Generate simple keyword metadata from a chunk.
    """
    tokens = tokenize(text)
    counts = Counter(tokens)
    most_common = [word for word, _ in counts.most_common(max_keywords)]

    return ", ".join(most_common)


def process_file_to_chunks(
    file_path: str,
    original_filename: str,
    chunk_size_words: int = 140,
    overlap_words: int = 30
) -> List[Dict]:
    """
    Extract text from a file and convert it into chunk records.
    """
    page_records = extract_document_text(file_path)

    all_chunks = []
    global_chunk_index = 0

    for page in page_records:
        page_number: Optional[int] = page.get("page_number")
        page_text = page.get("text", "")

        text_chunks = chunk_text(
            page_text,
            chunk_size_words=chunk_size_words,
            overlap_words=overlap_words
        )

        for chunk in text_chunks:
            keywords = generate_keywords(chunk)

            all_chunks.append(
                {
                    "source_filename": original_filename,
                    "page_number": page_number,
                    "chunk_index": global_chunk_index,
                    "chunk_text": chunk,
                    "keywords": keywords,
                }
            )

            global_chunk_index += 1

    return all_chunks
